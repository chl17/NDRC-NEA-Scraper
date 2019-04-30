# .docx
import docx
# .doc
import subprocess
# .pdf
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfdevice import PDFDevice
from pdfminer.layout import *
from pdfminer.converter import PDFPageAggregator
# .xls and .xlsx
import xlrd
# timeout
import func_timeout
from func_timeout import func_set_timeout
import time
# pdf2txt.py
import os


def cleanattachment(list):  # 清理附件地址，清除链接和 mailto
    attachment = []
    for link in list:
        if link.find('pdf') + 1 or link.find('doc') + 1 or link.find('xls') + 1 or link.find('xlsx') + 1 or \
                link.find('docx') + 1:
            attachment.append(link)
    return attachment


def determine_class0(link):
    if link.split('/')[3] == 'xwzx':
        return '新闻发布中心'
    elif link.split('/')[3] == 'fzgggz':
        return '发展改革工作'
    elif link.split('/')[2] == 'gjss.ndrc.gov.cn':
        return '高技术产业司'
    elif link.split('/')[2] == 'hzs.ndrc.gov.cn':
        return '资源节约和环境保护司'
    elif link.split('/')[2] == 'shs.ndrc.gov.cn':
        return '社会发展司'
    elif link.split('/')[2] == 'cjs.ndrc.gov.cn':
        return '财政金融司'
    elif link.split('/')[3] == 'zwfwzx':
        return '政务服务中心'
    elif link.split('/')[3] == 'zcfb':
        return '政策发布中心'
    elif link.split('/')[3] == 'govszyw':
        return '时政要闻'
    elif link.split('/')[3] == 'gzdt':
        return '委工作动态'
    elif link.split('/')[3] == 'dffgwdt':
        return '地方动态'
    elif link.split('/')[3] == 'jjxsfx':
        return '经济形势分析'
    else:
        return ''


def determine_class1(link):
    if link.split('/')[4] == 'xwfb':
        return '新闻发布'
    elif link.split('/')[4] == 'wszb':
        return '网上直播'
    elif link.split('/')[4] == 'fzgh':
        return '发展规划'
    elif link.split('/')[4] == 'hgjj':
        return '宏观经济'
    elif link.split('/')[4] == 'jjyx':
        return '经济运行'
    elif link.split('/')[4] == 'tzgg':
        return '体制改革'
    elif link.split('/')[4] == 'gdzctz':
        return '固定资产投资'
    elif link.split('/')[4] == 'wzly':
        return '外资利用'
    elif link.split('/')[4] == 'dqjj':
        return '地区经济'
    elif link.split('/')[4] == 'ncjj':
        return '农村经济'
    elif link.split('/')[4] == 'nyjt':
        return '基础产业'
    elif link.split('/')[4] == 'gyfz':
        return '产业发展'
    elif link.split('/')[3] == 'gjsgz':
        return '高技术工作'
    elif link.split('/')[3] == 'gzdtx':
        return '发展动态'
    elif link.split('/')[3] == 'ghzc':
        return '政策发布'
    elif link.split('/')[4] == 'hjbh':
        return '环境与资源'
    elif link.split('/')[3] == 'gzdt':
        return '社会发展工作'
    elif link.split('/')[3] == 'shfzdt':
        return '社会发展动态'
    elif link.split('/')[3] == 'zcyj':
        return '社会发展规划、政策与研究'
    elif link.split('/')[4] == 'jyysr':
        return '就业与收入'
    elif link.split('/')[4] == 'jjmy':
        return '经济贸易'
    elif link.split('/')[3] == 'shxytxjs':
        return '信用建设'
    elif link.split('/')[4] == 'jggl':
        return '价格管理'
    elif link.split('/')[4] == 'tztg':
        return '通知通告'
    elif link.split('/')[4] == 'xzxknew':
        return '行政许可'
    elif link.split('/')[4] == 'xzzq':
        return '下载专区'
    elif link.split('/')[4] == 'zfdj':
        return '政府定价'
    elif link.split('/')[4] == 'zcfbl':
        return '发展改革委令'
    elif link.split('/')[4] == 'gfxwj':
        return '规范性文件'
    elif link.split('/')[4] == 'zcfbgg':
        return '发展改革委令'
    elif link.split('/')[4] == 'zcfbl':
        return '公告'
    elif link.split('/')[4] == 'zcfbghwb':
        return '规划文本'
    elif link.split('/')[4] == 'zcfbtz':
        return '通知'
    elif link.split('/')[4] == 'jd':
        return '解读'
    elif link.split('/')[4] == 'zcfbqt':
        return '其他'
    else:
        return ''


def readattachment(file_paths):
    content_all = []
    for file_path in file_paths:
        file_path = '/Users/chenhaolin/PycharmProjects/SRT/发改委/NDRC/FILES' + file_path
        # .doc
        if file_path.endswith('.doc'):
            try:
                textlist = []
                word = file_path
                output = subprocess.check_output(["/Users/chenhaolin/PycharmProjects/SRT/发改委/NDRC/antiword", word])
                text = output.decode('utf-8')
                textlist.append(text.strip())
                content_all.append('\n'.join(textlist))
            except:
                content_all.append(file_path + ' 读取错误!')
        # .pdf
        if file_path.endswith('.pdf'):
            try:  # 防止读取 PDF 超时阻塞 使用 pdf2txt.py
                content = pdf2text(file_path)
                if content.strip():  # 排除无文字读取的 PDF
                    content_all.append(content)
            except func_timeout.exceptions.FunctionTimedOut:
                content_all.append(file_path + ' 读取超时!')
            '''
            try:  # 防止读取 PDF 超时阻塞 使用 PDFMiner
                content = readPDF(file_path)
                content_all.append(content)
            except func_timeout.exceptions.FunctionTimedOut:
                content_all.append(file_path + ' 读取超时!')
            '''
        # .docx
        if file_path.endswith('.docx'):
            try:
                doc = docx.Document(file_path)
                content = []
                for para in doc.paragraphs:
                    content.append(para.text)
                content_all.append('\n'.join(content))
            except:
                content_all.append(file_path + ' 读取错误!')
        # .xls and .xlsx
        if file_path.endswith('.xls') or file_path.endswith('.xlsx'):
            try:
                data = xlrd.open_workbook(file_path)
                table_names = data.sheet_names()
                content_list = []
                content_cleaned = []
                for table_name in table_names:
                    table = data.sheet_by_name(table_name)
                    col_count = table.ncols
                    for col in range(col_count):
                        content_list = content_list + table.col_values(col)
                for element in content_list:
                    if element != '' and (not ((type(element) == int) or (type(element) == float))):
                        content_cleaned.append(str(element))
                content_all.append(''.join(content_cleaned))
            except:
                content_all.append(file_path + ' 读取错误!')
    return ' '.join(content_all)


@func_set_timeout(5)  # 设定超时限制5s
def readPDF(file_path):
    fp = open(file_path, 'rb')
    # 来创建一个pdf文档分析器
    parser = PDFParser(fp)
    # 创建一个PDF文档对象存储文档结构
    document = PDFDocument(parser)
    # 检查文件是否允许文本提取
    if not document.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        # 创建一个PDF资源管理器对象来存储共赏资源
        rsrcmgr = PDFResourceManager()
        # 设定参数进行分析
        laparams = LAParams()
        # 创建一个PDF设备对象
        # device=PDFDevice(rsrcmgr)
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        # 创建一个PDF解释器对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        content = []
        content_cleaned = []
        # 处理每一页
        for page in PDFPage.create_pages(document):
            interpreter.process_page(page)
            # 接受该页面的LTPage对象
            layout = device.get_result()
            for x in layout:
                if (isinstance(x, LTTextBoxHorizontal)):
                    text_stripped = x.get_text().strip()
                    if text_stripped != '':
                        content.append(x.get_text().strip())
        for element in content:
            if element != '':
                content_cleaned.append(element)
        return '\n'.join(content_cleaned)


@func_set_timeout(5)  # 设定超时限制5s  http://www.cnblogs.com/hester/p/7641258.html
def pdf2text(file_path):
    # 路径修改
    cmd = 'python ' + '/Users/chenhaolin/PycharmProjects/SRT/发改委/NDRC/pdf2txt.py ' \
              + file_path
    output_text = os.popen(cmd)
    return output_text.read()

