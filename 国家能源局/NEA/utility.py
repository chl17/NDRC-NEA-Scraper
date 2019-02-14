# .docx
import docx

# .doc
import subprocess

# .pdf
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfdevice import PDFDevice
from pdfminer.layout import *
from pdfminer.converter import PDFPageAggregator

# .xls and .xlsx
import xlrd

# timeout
import func_timeout
from func_timeout import func_set_timeout
import time

# pdf2txt.py
import os


def cleanattachment(list):  # 清除链接和 mailto
    attachment = []
    for link in list:
        if link.find('pdf') + 1 or link.find('doc') + 1 or link.find('xls') + 1 or link.find('xlsx') + 1 or \
                link.find('docx') + 1:
            attachment.append(link)
    return attachment


def determine_class0(link):
    if link.split('/')[4].find('zxwj') >= 0:
        return '最新文件'
    elif link.split('/')[4].find('tz') >= 0:
        return '通知'
    elif link.split('/')[4].find('gg') >= 0:
        return '公告'
    elif link.split('/')[4].find('xmsp') >= 0:
        return '发展改革工作'
    elif link.split('/')[4].find('jd') >= 0:
        return '解读'
    elif link.split('/')[4].find('qt') >= 0:
        return '其他'
    elif link.split('/')[3].find('ldhd') >= 0:
        return '领导活动'
    elif link.split('/')[4].find('jwzdt') >= 0:
        return '局工作动态'
    elif link.split('/')[4].find('pcjg') >= 0:
        return '派出机构动态'
    elif link.split('/')[4].find('nyyw') >= 0:
        return '能源要闻'
    else:
        return ''


def readattachment(file_paths):
    content_all = []
    for file_path in file_paths:
        file_path = '/Users/chenhaolin/PycharmProjects/SRT/国家能源局/NEA/FILES' + file_path
        # .doc
        if file_path.endswith('.doc'):
            try:
                textlist = []
                word = file_path
                output = subprocess.check_output(["/Users/chenhaolin/PycharmProjects/SRT/国家能源局/NEA/antiword", word])
                text = output.decode('utf-8')
                textlist.append(text.strip())
                content_all.append('\n'.join(textlist))
            except:
                content_all.append(file_path + ' 读取错误!')
        # .pdf
        if file_path.endswith('.pdf'):
            try:  # 防止读取 PDF 超时阻塞 使用 pdf2txt.py
                content = pdf2text(file_path)
                if content.strip():  # 排除无文字读取的 PDF
                    content_all.append(content)
            except func_timeout.exceptions.FunctionTimedOut:
                content_all.append(file_path + ' 读取超时!')
            '''
            try:  # 防止读取 PDF 超时阻塞 使用 PDFMiner
                content = readPDF(file_path)
                content_all.append(content)
            except func_timeout.exceptions.FunctionTimedOut:
                content_all.append(file_path + ' 读取超时!')
            '''
        # .docx
        if file_path.endswith('.docx'):
            try:
                doc = docx.Document(file_path)
                content = []
                for para in doc.paragraphs:
                    content.append(para.text)
                content_all.append('\n'.join(content))
            except:
                content_all.append(file_path + ' 读取错误!')
        # .xls and .xlsx
        if file_path.endswith('.xls') or file_path.endswith('.xlsx'):
            try:
                data = xlrd.open_workbook(file_path)
                table_names = data.sheet_names()
                content_list = []
                content_cleaned = []
                for table_name in table_names:
                    table = data.sheet_by_name(table_name)
                    col_count = table.ncols
                    for col in range(col_count):
                        content_list = content_list + table.col_values(col)
                for element in content_list:
                    if element != '' and (not ((type(element) == int) or (type(element) == float))):
                        content_cleaned.append(str(element))
                content_all.append(''.join(content_cleaned))
            except:
                content_all.append(file_path + ' 读取错误!')
    return ' '.join(content_all)


@func_set_timeout(5)  # 设定超时限制5s
def readPDF(file_path):
    fp = open(file_path, 'rb')
    # 来创建一个pdf文档分析器
    parser = PDFParser(fp)
    # 创建一个PDF文档对象存储文档结构
    document = PDFDocument(parser)
    # 检查文件是否允许文本提取
    if not document.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        # 创建一个PDF资源管理器对象来存储共赏资源
        rsrcmgr = PDFResourceManager()
        # 设定参数进行分析
        laparams = LAParams()
        # 创建一个PDF设备对象
        # device=PDFDevice(rsrcmgr)
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        # 创建一个PDF解释器对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        content = []
        content_cleaned = []
        # 处理每一页
        for page in PDFPage.create_pages(document):
            interpreter.process_page(page)
            # 接受该页面的LTPage对象
            layout = device.get_result()
            for x in layout:
                if (isinstance(x, LTTextBoxHorizontal)):
                    text_stripped = x.get_text().strip()
                    if text_stripped != '':
                        content.append(x.get_text().strip())
        for element in content:
            if element != '':
                content_cleaned.append(element)
        return '\n'.join(content_cleaned)


@func_set_timeout(5)  # 设定超时限制5s
def pdf2text(file_path):
    cmd = 'python3 ' + '/Users/chenhaolin/PycharmProjects/SRT/国家能源局/NEA/pdf2txt.py ' \
              + file_path
    output_text = os.popen(cmd)
    return output_text.read()


"""
file_path = ['/Users/chenhaolin/PycharmProjects/SRT/发改委/NDRC/FILES/test2.pdf']
content = extract_text(file_path)
print(str(content).strip())
print('\n')
print(type(content))

"""

"""
command = 'python3 ' + '/Users/chenhaolin/PycharmProjects/SRT/发改委/NDRC/pdf2txt.py ' \
          + '/Users/chenhaolin/PycharmProjects/SRT/发改委/NDRC/FILES/test7.pdf'
output = os.popen(command)
print(output.read())
"""

'''
paths = ['/test1.pdf',
         '/test2.pdf',
         '/test3.pdf',
         '/test4.pdf',
         '/test5.pdf',
         '/test6.pdf',
         '/test7.pdf',
         '/test8.pdf',
         '/test9.pdf',
         '/test10.pdf']
content_para = readattachment(paths)
for content0 in content_para:
    # print(content + '\n' + '*' * 50)
    print(content0)
    print('\n' + '*' * 50)
'''

"""

paths = ['/test.doc',
         '/test.docx',
         '/test.pdf',
         '/test_xls.xls',
         '/test_xlsx.xlsx',]
content_para = readattachment(paths)
for content0 in content_para:
    # print(content + '\n' + '*' * 50)
    print(content0)
    print('\n' + '*' * 50)
"""

"""

file_path = ['/test.pdf']
data = readattachment(file_path)
print(data)
"""

"""
file_path = '/Users/chenhaolin/PycharmProjects/SRT/发改委/NDRC/FILES/test.pdf'
print(readPDF(file_path))
print(type(readPDF(file_path)))
"""

"""
file_path = '/Users/chenhaolin/PycharmProjects/SRT/发改委/NDRC/FILES/2017-08-10_十三五”水利科技创新规划/W020170810372162276685.doc'
if file_path.find('doc') > 0 and (not (file_path.find('docx') > 0)):
    file_path = file_path
    print('1')
    print(readdoc(file_path))
"""

"""
word = "test.doc"
output = subprocess.check_output(["/Users/chenhaolin/PycharmProjects/SRT/发改委/NDRC/antiword", word])
a = output.decode('utf-8')
print(a)
"""





"""
item = ['0', '1', '2', '3', '4', '5', '6', '7']
try:
    if item[9]:
        print('s')
except IndexError:
    pass


list0 = ['mailto:fgwwhkl@163.com', './W020180717396249402008.docx', 'http://baidu.com']
list1 = cleanattachment(list0)
print(list1)

link0 = 'mailto:fgwwhkl@163.com'
link1 = './W020180717396249402008.docx'
print(link0.find('mailto'))
print(link1.find('mailto'))
"""